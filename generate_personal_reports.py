# -*- coding: utf-8 -*-
"""
生成4个个人文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table_with_header(doc, headers, rows):
    """添加带表头的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], 'D9EAD3')
        header_cells[i].paragraphs[0].runs[0].bold = True
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
    return table

def create_zhou_document():
    """周志文文档 - 组长：系统需求、实施方案、业务系统结构、改进建议"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)

    # 标题
    title = doc.add_heading('', 0)
    run = title.add_run('上海杉达学院')
    run.font.size = Pt(18)
    run.bold = True
    title2 = doc.add_heading('', 0)
    run2 = title2.add_run('教学一体化平台升级改造中期自查报告')
    run2.font.size = Pt(22)
    run2.bold = True
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('【负责人：周志文】').bold = True
    p.add_run('（组长）')

    doc.add_paragraph()

    # 1 系统需求
    add_heading(doc, '1  系统需求', 1)
    add_heading(doc, '1.1  系统需求说明', 2)
    doc.add_paragraph(
        '本项目旨在将原有的Spring MVC + JSP + jQuery实验教学管理系统（Labex）'
        '重构为现代化的Vue 3 + Spring Boot 3前后端分离架构。系统需要支持教师和学生的'
        '全流程实验教学管理，包括班级管理、学生管理、实验管理、成绩管理等核心功能。'
        '重构过程中需完全保留原有数据库结构，确保数据兼容性。'
    )

    add_heading(doc, '1.2  系统实现情况', 2)
    headers = ['序号', '模块名称', '功能概述', '承诺人', '完成日期', '完成比例']
    rows = [
        ['1', '系统架构设计', '完成前后端分离架构设计，定义清晰的接口规范', '周志文', '2026-04-15', '100%'],
        ['2', '后端核心模块', '完成Spring Boot 3后端框架搭建及核心业务逻辑实现', '周志文', '2026-04-16', '100%'],
        ['7', '安全认证模块', '完成JWT认证、权限控制', '周志文', '2026-04-14', '100%'],
    ]
    add_table_with_header(doc, headers, rows)

    add_heading(doc, '1.3  详细实现情况及存在差异', 2)
    doc.add_paragraph(
        '【已完成】系统架构采用Vue 3 + Spring Boot 3前后端分离架构，'
        '后端使用Spring Security + JWT实现认证授权，前端使用Pinia进行状态管理。'
        '原有20+张数据库表结构完全保留，新增12个索引优化查询性能。'
    )
    doc.add_paragraph(
        '【差异说明】原计划使用Shiro进行权限控制，实际采用Spring Security 6.x，'
        '因Spring Security与Spring Boot 3集成更紧密，安全性更高。'
        '原计划使用Session认证，实际采用JWT Token认证，支持无状态分布式部署。'
    )

    doc.add_page_break()

    # 2 实施方案
    add_heading(doc, '2  实施方案', 1)
    add_heading(doc, '2.1  技术路线选择', 2)
    headers = ['序号', '项目', '计划', '实际', '差异']
    rows = [
        ['1', '后端框架', 'Spring MVC 4', 'Spring Boot 3', '架构升级'],
        ['2', '前端框架', 'jQuery + JSP', 'Vue 3 + Vite', '完全重构'],
        ['3', '持久层', 'MyBatis', 'MyBatis-Plus', '自动CRUD更高效'],
        ['4', '安全框架', 'Shiro', 'Spring Security + JWT', '更完善的权限控制'],
        ['5', '认证方式', 'Session', 'JWT Token', '无状态认证'],
    ]
    add_table_with_header(doc, headers, rows)

    add_heading(doc, '2.2  技术路线选择及存在差异原因', 2)
    doc.add_paragraph(
        '1. Spring Boot 3相比Spring MVC 4：提供更简化的配置、自动装配机制，'
        '与Spring Security 6.x兼容性更好，支持JDK 17+新特性。'
    )
    doc.add_paragraph(
        '2. Vue 3相比jQuery：组件化开发模式、响应式数据绑定、虚拟DOM，'
        '大幅提升开发效率和代码可维护性。'
    )
    doc.add_paragraph(
        '3. JWT Token：支持无状态认证，便于分布式部署和水平扩展。'
    )

    doc.add_page_break()

    # 4 业务系统结构及代码模式
    add_heading(doc, '4  业务系统结构及代码模式', 1)
    add_heading(doc, '4.1  业务系统结构', 2)
    headers = ['序号', '业务系统结构', '说明', '原系统', '新系统']
    rows = [
        ['1', '整体架构', '前后端分离', '前后端耦合（JSP）', '完全分离'],
        ['2', '后端模块', 'Controller-Service-Mapper', 'Service-Action-Model', '标准三层架构'],
        ['3', '前端模块', 'Vue组件化', '页面模板', '组件化+状态管理'],
        ['4', '认证方式', 'JWT Token无状态认证', 'Session认证', '无状态+RefreshToken'],
    ]
    add_table_with_header(doc, headers, rows)

    add_heading(doc, '4.2  代码模式', 2)
    headers = ['序号', '代码模式', '说明', '原系统', '新系统']
    rows = [
        ['1', '后端分层', 'Controller-Service-Mapper三层', 'Action-Service-Dao多层', '标准三层更清晰'],
        ['2', 'ORM框架', 'MyBatis-Plus增强', 'MyBatis', '自动CRUD更高效'],
        ['3', '前端组件', 'Vue 3 Composition API', '无组件化', '完全组件化'],
        ['4', '状态管理', 'Pinia集中管理', '无', '集中式状态管理'],
    ]
    add_table_with_header(doc, headers, rows)
    doc.add_paragraph(
        '【代码模式说明】后端采用标准三层架构（Controller-Service-Mapper），'
        '使用MyBatis-Plus实现数据访问自动化；前端采用Vue 3 Composition API进行组件化开发。'
    )

    doc.add_page_break()

    # 6 改进建议
    add_heading(doc, '6  改进建议', 1)
    add_heading(doc, '6.1  实现改进建议', 2)
    doc.add_paragraph('1. 【代码生成】引入MyBatis-Plus Generator自动生成基础CRUD代码。')
    doc.add_paragraph('2. 【接口文档】引入SpringDoc OpenAPI 3自动生成API文档。')
    doc.add_paragraph('3. 【参数校验】增强参数校验，使用@Validated和自定义校验注解。')
    doc.add_paragraph('4. 【缓存优化】针对热点数据引入Redis二级缓存。')
    doc.add_paragraph('5. 【异步处理】将邮件通知、成绩推送等非核心功能异步化。')

    doc.save(r'D:\workfordasan\docs\周志文_中期自查报告.docx')
    print("周志文文档已保存")

def create_yufei_document():
    """余飞杨文档 - 数据库部分"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)

    title = doc.add_heading('', 0)
    run = title.add_run('上海杉达学院')
    run.font.size = Pt(18)
    run.bold = True
    title2 = doc.add_heading('', 0)
    run2 = title2.add_run('教学一体化平台升级改造中期自查报告')
    run2.font.size = Pt(22)
    run2.bold = True
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('【负责人：余飞杨】')

    doc.add_paragraph()

    # 3 数据库
    add_heading(doc, '3  数据库', 1)

    add_heading(doc, '3.1  数据表', 2)
    doc.add_paragraph('系统共保留和优化了以下核心数据表：')
    headers = ['序号', '数据表名', '说明', '原数据库', '新数据库']
    rows = [
        ['1', 't_teacher', '教师/管理员表', '保留', '保留'],
        ['2', 't_student', '学生表', '保留', '保留'],
        ['3', 't_clazz', '班级表', '保留', '增强（新增teacher_id）'],
        ['4', 't_experiment', '实验表', '保留', '保留'],
        ['5', 't_experiment_item', '实验题目表', '保留', '保留'],
        ['6', 't_student_item', '学生答题表', '保留', '保留'],
        ['7', 't_score', '成绩表', '保留', '保留'],
        ['8', 't_lecture', '讲义表', '保留', '保留'],
        ['9', 't_question', '题库表', '新增', '保留'],
        ['10', 't_exam', '考试表', '新增', '保留'],
        ['11', 't_paper', '试卷表', '新增', '保留'],
        ['12', 't_paper_question', '试卷题目关联表', '新增', '保留'],
        ['13', 't_student_answer', '学生答案记录表', '新增', '保留'],
        ['14', 't_sys_log', '系统日志表', '保留', '保留'],
        ['15', 't_student_log', '学生日志表', '保留', '保留'],
    ]
    add_table_with_header(doc, headers, rows)

    add_heading(doc, '3.2  视图及存储过程', 2)
    doc.add_paragraph('数据库视图及存储过程情况：')
    headers = ['序号', '视图/存储过程', '说明', '原数据库', '新数据库']
    rows = [
        ['1', '无', '未使用视图', '未使用', '未使用'],
        ['2', '无', '未使用存储过程', '未使用', '未使用'],
        ['3', '索引优化', '新增12个索引优化查询性能', '部分索引', '新增12个索引'],
    ]
    add_table_with_header(doc, headers, rows)

    add_heading(doc, '3.3  数据库差异及存在差异原因', 2)
    doc.add_paragraph('【差异说明】数据库结构保持高度一致，主要差异在于：')
    doc.add_paragraph('1. 新增字段：t_clazz表新增teacher_id字段，关联所属教师，便于查询和管理。')
    doc.add_paragraph('2. 新增表：为支持新的考试功能，新增t_exam、t_paper、t_paper_question等表。')
    doc.add_paragraph('3. 索引优化：在关键查询字段上新增12个索引，提升查询性能。')
    doc.add_paragraph('4. 密码加密：密码存储从MD5升级为BCrypt加密，更安全。')

    doc.add_page_break()

    # 6.2 运维改进建议
    add_heading(doc, '6  改进建议', 1)
    add_heading(doc, '6.2  运维改进建议', 2)
    doc.add_paragraph('运维方面改进建议：')
    doc.add_paragraph('1. 【监控告警】引入Spring Boot Actuator + Prometheus监控指标。')
    doc.add_paragraph('2. 【日志规范】统一日志格式，引入ELK日志收集分析。')
    doc.add_paragraph('3. 【配置中心】引入Apollo或Nacos实现配置中心化管理。')
    doc.add_paragraph('4. 【容器编排】完善Kubernetes部署配置，支持弹性伸缩。')

    doc.save(r'D:\workfordasan\docs\余飞杨_中期自查报告.docx')
    print("余飞杨文档已保存")

def create_yang_document():
    """杨宇润文档 - 学生模块、运维业务功能"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)

    title = doc.add_heading('', 0)
    run = title.add_run('上海杉达学院')
    run.font.size = Pt(18)
    run.bold = True
    title2 = doc.add_heading('', 0)
    run2 = title2.add_run('教学一体化平台升级改造中期自查报告')
    run2.font.size = Pt(22)
    run2.bold = True
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('【负责人：杨宇润】')

    doc.add_paragraph()

    # 5 详细业务功能描述
    add_heading(doc, '5  详细业务功能描述', 1)

    add_heading(doc, '5.1  实现业务功能', 2)
    doc.add_paragraph(
        '原有业务功能已在重构后系统中完整实现，主要包括：'
    )
    doc.add_paragraph(
        '【学生模块功能】（杨宇润负责实现）'
    )
    doc.add_paragraph(
        '1. 实验列表：查看可用实验、进入答题'
    )
    doc.add_paragraph(
        '2. 实验答题：题目导航、在线答题、自动保存、提交答卷'
    )
    doc.add_paragraph(
        '3. 讲义学习：浏览和下载讲义'
    )
    doc.add_paragraph(
        '4. 个人中心：修改密码、查看成绩'
    )
    doc.add_paragraph(
        '【前端实现】学生端使用Vue 3 + Element Plus实现，包含：'
        'Experiment.vue（实验列表）、ExperimentDetail.vue（实验答题页）、'
        'Lecture.vue（讲义列表）、Profile.vue（个人中心）。'
        '实现了左侧题目导航实时高亮、答案自动保存（防抖500ms）、'
        '本地localStorage备份、进度条显示等功能。'
    )

    add_heading(doc, '5.2  运维业务功能', 2)
    doc.add_paragraph('运维相关业务功能已实现：')
    doc.add_paragraph(
        '【系统监控】操作日志记录（t_sys_log）、学生日志记录（t_student_log）、'
        '答题日志跟踪（t_student_item_log）。'
    )
    doc.add_paragraph(
        '【数据管理】数据备份支持、批量导入/导出（CSV格式）、Excel成绩导出。'
    )
    doc.add_paragraph(
        '【配置管理】系统参数可配置（t_sys_config）、题目类型可管理。'
    )
    doc.add_paragraph(
        '【前端实现】使用Vue 3 Composition API实现日志展示和数据导出功能。'
    )

    add_heading(doc, '5.3  其他业务功能', 2)
    doc.add_paragraph(
        '【认证授权】用户登录（JWT Token）、退出登录、Token自动刷新、用户信息获取。'
    )
    doc.add_paragraph(
        '【前端实现】Login.vue实现登录页（浅绿渐变背景、动画效果）。'
    )

    doc.add_page_break()

    # 6.3 测试建议
    add_heading(doc, '6  改进建议', 1)
    add_heading(doc, '6.3  测试建议', 2)
    doc.add_paragraph('测试方面改进建议：')
    doc.add_paragraph('1. 【单元测试】为Service层编写单元测试，使用JUnit 5 + Mockito。')
    doc.add_paragraph('2. 【集成测试】编写Controller层集成测试，验证接口功能。')
    doc.add_paragraph('3. 【前端测试】引入Vitest进行前端单元测试，Playwright进行E2E测试。')
    doc.add_paragraph('4. 【性能测试】使用JMeter进行接口性能测试，验证并发能力。')
    doc.add_paragraph('5. 【安全测试】进行SQL注入、XSS等安全测试。')

    doc.save(r'D:\workfordasan\docs\杨宇润_中期自查报告.docx')
    print("杨宇润文档已保存")

def create_yufeng_document():
    """于峰文档 - 教师模块、其他业务功能"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)

    title = doc.add_heading('', 0)
    run = title.add_run('上海杉达学院')
    run.font.size = Pt(18)
    run.bold = True
    title2 = doc.add_heading('', 0)
    run2 = title2.add_run('教学一体化平台升级改造中期自查报告')
    run2.font.size = Pt(22)
    run2.bold = True
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('【负责人：于峰】')

    doc.add_paragraph()

    # 5 详细业务功能描述
    add_heading(doc, '5  详细业务功能描述', 1)

    add_heading(doc, '5.1  实现业务功能', 2)
    doc.add_paragraph(
        '原有业务功能已在重构后系统中完整实现，主要包括：'
    )
    doc.add_paragraph(
        '【教师模块功能】（于峰负责实现）'
    )
    doc.add_paragraph(
        '1. 班级管理：增删改查班级、启用/禁用、IP限制'
    )
    doc.add_paragraph(
        '2. 学生管理：学生CRUD、批量导入（CSV）、重置密码/IP'
    )
    doc.add_paragraph(
        '3. 实验管理：创建/编辑实验、题目管理、标准答案设置'
    )
    doc.add_paragraph(
        '4. 讲义管理：文件上传、讲义列表'
    )
    doc.add_paragraph(
        '5. 成绩管理：查看学生答案、在线评分、批量评分、导出Excel'
    )
    doc.add_paragraph(
        '【后端实现】教师模块使用Spring Boot 3 + MyBatis-Plus实现，包含：'
        'ClazzController（班级管理）、StudentController（学生管理）、'
        'ExperimentController（实验管理）、LectureController（讲义管理）、'
        'ScoreController（成绩管理）、HomeworkController（作业管理）。'
    )

    add_heading(doc, '5.3  其他业务功能', 2)
    doc.add_paragraph(
        '【班级管理】班级启用/禁用、IP限制设置。'
    )
    doc.add_paragraph(
        '【学生管理】学生账号启用/禁用、IP地址绑定、密码错误次数限制。'
    )
    doc.add_paragraph(
        '【考试功能】新增Exam模块，支持在线考试功能（ExamController、ExamService）。'
    )
    doc.add_paragraph(
        '【评分功能】新增QuestionTestCase模块，支持编程题自动评分（QuestionTestCaseService）。'
    )

    add_heading(doc, '6.3  测试建议', 2)
    doc.add_paragraph('测试方面改进建议：')
    doc.add_paragraph('1. 【成绩模块测试】重点测试评分、批量评分、成绩导出功能。')
    doc.add_paragraph('2. 【考试模块测试】测试考试创建、答题、提交、评分完整流程。')
    doc.add_paragraph('3. 【并发测试】验证多学生同时答题时的数据一致性。')

    doc.save(r'D:\workfordasan\docs\于峰_中期自查报告.docx')
    print("于峰文档已保存")

if __name__ == '__main__':
    create_zhou_document()
    create_yufei_document()
    create_yang_document()
    create_yufeng_document()
    print("所有个人文档生成完成！")